"""
Azure Document Intelligence

Parse PDFs, Word docs, spreadsheets, and text files from Discord attachments.
Answer questions, summarize, extract data, and compare versions.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any


@dataclass
class DocumentContent:
    """Parsed document content."""
    text: str
    pages: list[str] = None
    tables: list[list[str]] = None
    metadata: dict[str, Any] = None
    source: str = ""

    def __post_init__(self):
        if self.pages is None:
            self.pages = []
        if self.tables is None:
            self.tables = []
        if self.metadata is None:
            self.metadata = {}


class DocumentIntelligence:
    """
    Document parsing and analysis system.

    Usage:
        di = DocumentIntelligence()
        doc = di.parse_attachment(bytes, "report.pdf")
        summary = di.summarize(doc)
        answer = di.answer_question(doc, "What was the revenue?")
    """

    def __init__(self):
        self._parsers = {}
        self._init_parsers()

    def _init_parsers(self):
        """Detect available parsers."""
        try:
            import PyPDF2
            self._parsers["pdf"] = PyPDF2
        except ImportError:
            pass

        try:
            import docx
            self._parsers["docx"] = docx
        except ImportError:
            pass

        try:
            import openpyxl
            self._parsers["xlsx"] = openpyxl
        except ImportError:
            pass

    # ------------------------------------------------------------------
    # Parsing
    # ------------------------------------------------------------------

    def parse_attachment(self, data: bytes, filename: str) -> DocumentContent:
        """Parse a document from bytes."""
        ext = filename.split(".")[-1].lower() if "." in filename else ""

        if ext == "pdf" and "pdf" in self._parsers:
            return self._parse_pdf(data, filename)
        elif ext in ("docx", "doc") and "docx" in self._parsers:
            return self._parse_docx(data, filename)
        elif ext in ("xlsx", "xls", "csv") and "xlsx" in self._parsers:
            return self._parse_xlsx(data, filename)
        elif ext in ("txt", "md", "py", "js", "json", "yaml", "yml"):
            return self._parse_text(data, filename)

        return DocumentContent(
            text=f"[Unsupported file type: {ext}. Install PyPDF2, python-docx, or openpyxl for better support.]",
            source=filename,
        )

    def _parse_pdf(self, data: bytes, filename: str) -> DocumentContent:
        """Parse a PDF file."""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            full_text = "\n\n".join(pages)
            return DocumentContent(
                text=full_text[:10000],  # Cap at 10k chars
                pages=pages[:20],
                metadata={"pages": len(reader.pages), "type": "pdf"},
                source=filename,
            )
        except Exception as e:
            return DocumentContent(text=f"[PDF parsing error: {e}]", source=filename)

    def _parse_docx(self, data: bytes, filename: str) -> DocumentContent:
        """Parse a Word document."""
        try:
            import docx
            doc = docx.Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            return DocumentContent(
                text=full_text[:10000],
                pages=paragraphs[:100],
                metadata={"paragraphs": len(paragraphs), "type": "docx"},
                source=filename,
            )
        except Exception as e:
            return DocumentContent(text=f"[DOCX parsing error: {e}]", source=filename)

    def _parse_xlsx(self, data: bytes, filename: str) -> DocumentContent:
        """Parse an Excel spreadsheet."""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data))
            sheets = []
            tables = []
            for sheet_name in wb.sheetnames[:3]:  # Limit to first 3 sheets
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(max_row=min(sheet.max_row, 50)):
                    row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                    rows.append(row_data)
                if rows:
                    tables.append(rows)
                    sheets.append(f"**Sheet: {sheet_name}**\n" + "\n".join([" | ".join(r[:10]) for r in rows[:5]]))
            return DocumentContent(
                text="\n\n".join(sheets)[:10000],
                tables=tables,
                metadata={"sheets": len(wb.sheetnames), "type": "xlsx"},
                source=filename,
            )
        except Exception as e:
            return DocumentContent(text=f"[XLSX parsing error: {e}]", source=filename)

    def _parse_text(self, data: bytes, filename: str) -> DocumentContent:
        """Parse a plain text file."""
        try:
            text = data.decode("utf-8", errors="replace")
            return DocumentContent(
                text=text[:10000],
                pages=text.split("\n\n")[:50],
                metadata={"chars": len(text), "type": "text"},
                source=filename,
            )
        except Exception as e:
            return DocumentContent(text=f"[Text parsing error: {e}]", source=filename)

    # ------------------------------------------------------------------
    # Analysis
    # ------------------------------------------------------------------

    def summarize(self, doc: DocumentContent, max_sentences: int = 5) -> str:
        """Extractive summarization using simple sentence scoring."""
        text = doc.text
        if not text or len(text) < 200:
            return text

        sentences = text.replace("!", ".").replace("?", ".").split(".")
        sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

        if len(sentences) <= max_sentences:
            return text[:1000]

        # Score sentences by word frequency
        word_freq = {}
        for s in sentences:
            for word in s.lower().split():
                word_freq[word] = word_freq.get(word, 0) + 1

        sentence_scores = []
        for s in sentences:
            score = sum(word_freq.get(w.lower(), 0) for w in s.split())
            sentence_scores.append((score, s))

        sentence_scores.sort(reverse=True)
        top = [s for _, s in sentence_scores[:max_sentences]]
        return "\n".join(f"• {s}" for s in top)

    def answer_question(self, doc: DocumentContent, question: str) -> str:
        """Simple question answering via keyword matching."""
        doc.text.lower()
        q_words = [w for w in question.lower().split() if len(w) > 3]

        # Find relevant paragraphs
        paragraphs = doc.text.split("\n\n")
        best_para = ""
        best_score = 0

        for para in paragraphs:
            para_lower = para.lower()
            score = sum(1 for w in q_words if w in para_lower)
            if score > best_score:
                best_score = score
                best_para = para

        if best_para and best_score > 0:
            return f"Based on the document:\n\n{best_para[:500]}"

        return "I couldn't find a specific answer in the document. Try rephrasing your question."

    def compare_documents(self, doc1: DocumentContent, doc2: DocumentContent) -> str:
        """Compare two documents and list differences."""
        text1 = doc1.text
        text2 = doc2.text

        if not text1 or not text2:
            return "Cannot compare: one or both documents are empty."

        # Simple line-by-line comparison
        lines1 = set(text1.split("\n"))
        lines2 = set(text2.split("\n"))

        only_in_1 = list(lines1 - lines2)[:10]
        only_in_2 = list(lines2 - lines1)[:10]

        result = ["**Document Comparison**"]
        if only_in_1:
            result.append(f"\nOnly in {doc1.source}:")
            for line in only_in_1:
                result.append(f"- {line[:80]}")
        if only_in_2:
            result.append(f"\nOnly in {doc2.source}:")
            for line in only_in_2:
                result.append(f"- {line[:80]}")
        if not only_in_1 and not only_in_2:
            result.append("\nDocuments are very similar or identical.")

        return "\n".join(result)
